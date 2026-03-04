"""
导出服务
支持多种格式导出
"""


import io
import json
import hashlib
import zipfile
from typing import List
from datetime import datetime

from src.models import TestCase
from src.utils.validators import cell_text, escape_html
from src.config import TABLE_HEADERS, get_type_tag_style


class ExportService:
    """导出服务"""
    
    def to_markdown(self, cases: List[TestCase]) -> str:
        """导出为Markdown"""
        if not cases:
            return ""
        
        lines = [
            "| " + " | ".join(TABLE_HEADERS) + " |",
            "| " + " | ".join(["---"] * len(TABLE_HEADERS)) + " |"
        ]
        
        for case in cases:
            row = [
                cell_text(case.case_id),
                cell_text(case.case_name),
                cell_text(case.module),
                cell_text(case.test_type.value),
                cell_text(case.precondition),
                cell_text(case.steps),
                cell_text(case.test_data),
                cell_text(case.expected),
                cell_text(case.priority.value),
            ]
            lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(lines)
    
    def to_excel(self, cases: List[TestCase]) -> bytes:
        """导出为Excel"""
        import pandas as pd
        from openpyxl import load_workbook
        from openpyxl.styles import PatternFill
        
        rows = []
        for case in cases:
            rows.append([
                cell_text(case.case_id),
                cell_text(case.case_name),
                cell_text(case.module),
                cell_text(case.test_type.value),
                cell_text(case.precondition),
                cell_text(case.steps),
                cell_text(case.test_data),
                cell_text(case.expected),
                cell_text(case.priority.value),
            ])
        
        df = pd.DataFrame(rows, columns=TABLE_HEADERS)
        buf = io.BytesIO()
        df.to_excel(buf, index=False, engine="openpyxl")
        buf.seek(0)
        
        # 应用样式
        wb = load_workbook(buf)
        ws = wb.active
        
        fills = {
            "高": PatternFill(start_color="FFCCCB", fill_type="solid"),
            "中": PatternFill(start_color="FFFFCC", fill_type="solid"),
            "低": PatternFill(start_color="90EE90", fill_type="solid"),
        }
        
        for row in range(2, ws.max_row + 1):
            pri = ws.cell(row=row, column=9).value
            if pri in fills:
                for col in range(1, 10):
                    ws.cell(row=row, column=col).fill = fills[pri]
        
        # 添加筛选
        if ws.max_row >= 1:
            ws.auto_filter.ref = f"A1:I{ws.max_row}"
        
        out = io.BytesIO()
        wb.save(out)
        out.seek(0)
        return out.read()
    
    def to_word(self, cases: List[TestCase]) -> bytes:
        """导出为Word"""
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading("测试用例", 0)
        doc.add_paragraph("以下为根据需求生成的测试用例列表，可直接用于执行。")
        doc.add_paragraph()
        doc.add_heading("目录", level=1)
        
        for case in cases:
            doc.add_paragraph(f"{case.case_id} - {case.case_name}", style="List Bullet")
        
        doc.add_page_break()
        
        for i, case in enumerate(cases):
            if i > 0:
                doc.add_page_break()
            
            doc.add_heading(f"{case.case_id} - {case.case_name}", level=1)
            doc.add_paragraph(f"所属模块：{case.module}")
            doc.add_paragraph(f"测试类型：{case.test_type.value} | 优先级：{case.priority.value}")
            doc.add_paragraph(f"前置条件：{case.precondition}")
            doc.add_paragraph(f"测试步骤：{case.steps}")
            doc.add_paragraph(f"测试数据：{case.test_data}")
            doc.add_paragraph(f"预期结果：{case.expected}")
        
        # 设置页眉页脚
        for s in doc.sections:
            s.header.is_linked_to_previous = False
            s.footer.is_linked_to_previous = False
            header = s.header.paragraphs[0] if s.header.paragraphs else s.header.add_paragraph()
            header.text = "TestGen AI - 测试用例"
            footer = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
            footer.text = "测试用例文档"
        
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    
    def to_pdf(self, cases: List[TestCase]) -> bytes:
        """导出为PDF"""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=12 * mm,
            rightMargin=12 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
        )
        
        # 尝试加载中文字体
        font_name = "Helvetica"
        font_paths = [
            "~/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "C:/Windows/Fonts/msyh.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        ]
        
        import os
        for fp in font_paths:
            full_path = os.path.expanduser(fp)
            if os.path.exists(full_path):
                try:
                    font_name = "PDFChinese"
                    pdfmetrics.registerFont(TTFont(font_name, full_path))
                    break
                except Exception:
                    continue
        
        # 表头和数据
        headers = ["编号", "名称", "模块", "类型", "前置", "步骤", "数据", "预期", "优先级"]
        col_widths = [22*mm, 28*mm, 22*mm, 18*mm, 22*mm, 35*mm, 22*mm, 35*mm, 18*mm]
        
        data = [headers]
        for case in cases:
            data.append([
                cell_text(case.case_id)[:12],
                cell_text(case.case_name)[:20],
                cell_text(case.module)[:12],
                cell_text(case.test_type.value)[:8],
                cell_text(case.precondition)[:14],
                cell_text(case.steps)[:24],
                cell_text(case.test_data)[:12],
                cell_text(case.expected)[:24],
                cell_text(case.priority.value)[:4],
            ])
        
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#667eea")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ]))
        
        doc.build([t])
        buf.seek(0)
        return buf.read()
    
    def to_xmind(self, cases: List[TestCase]) -> bytes:
        """导出为XMind"""
        # 按模块分组
        by_module = {}
        for case in cases:
            mod = case.module or "未分类"
            if mod not in by_module:
                by_module[mod] = []
            by_module[mod].append(case)
        
        counter = [0]
        
        def make_topic(title, children=None):
            counter[0] += 1
            tid = f"topic{counter[0]}"
            node = {"id": tid, "title": title[:2000]}
            if children:
                node["children"] = {"attached": children}
            return node
        
        case_topics = []
        for mod, mod_cases in by_module.items():
            mod_children = []
            for c in mod_cases:
                cid = cell_text(c.case_id)
                name = cell_text(c.case_name)
                
                sub = []
                if c.precondition:
                    sub.append(make_topic(f"前置：{cell_text(c.precondition)}"))
                if c.steps:
                    sub.append(make_topic(f"步骤：{cell_text(c.steps)}"))
                if c.test_data:
                    sub.append(make_topic(f"数据：{cell_text(c.test_data)}"))
                if c.expected:
                    sub.append(make_topic(f"预期：{cell_text(c.expected)}"))
                if c.priority:
                    sub.append(make_topic(f"优先级：{c.priority.value}"))
                
                mod_children.append(make_topic(f"{cid} - {name}", sub if sub else None))
            
            case_topics.append(make_topic(mod, mod_children))
        
        counter[0] += 1
        root = {
            "id": f"root{counter[0]}",
            "title": "测试用例",
            "children": {"attached": case_topics}
        }
        
        sheet = {"id": "sheet-1", "title": "测试用例", "rootTopic": root}
        content = [sheet]
        content_bytes = json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")
        
        now = int(datetime.now().timestamp() * 1000)
        metadata = {
            "creator": {"name": "TestGen AI", "version": "8.0"},
            "created": {"timestamp": now},
            "modified": {"timestamp": now},
        }
        metadata_bytes = json.dumps(metadata, ensure_ascii=False).encode("utf-8")
        
        def checksum(b):
            return hashlib.sha256(b).hexdigest()
        
        manifest = {
            "file-entries": {
                "content.json": {"checksum": checksum(content_bytes)},
                "metadata.json": {"checksum": checksum(metadata_bytes)},
            }
        }
        
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
            z.writestr("content.json", content_bytes.decode("utf-8"))
            z.writestr("metadata.json", metadata_bytes.decode("utf-8"))
        
        buf.seek(0)
        return buf.read()
    
    def to_opml(self, cases: List[TestCase]) -> bytes:
        """导出为OPML"""
        from xml.sax.saxutils import escape
        
        lines = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<opml version="2.0">',
            "<head><title>测试用例</title></head>",
            "<body>",
            '<outline text="测试用例">',
        ]
        
        # 按模块分组
        by_module = {}
        for case in cases:
            mod = case.module or "未分类"
            if mod not in by_module:
                by_module[mod] = []
            by_module[mod].append(case)
        
        for mod, mod_cases in by_module.items():
            lines.append(f'<outline text="{escape(mod)}">')
            for c in mod_cases:
                cid = cell_text(c.case_id)
                name = cell_text(c.case_name)
                title = f"{cid} - {name}"
                lines.append(f'<outline text="{escape(title)}">')
                
                if c.precondition:
                    lines.append(f'<outline text="前置：{escape(cell_text(c.precondition))}"/>')
                if c.steps:
                    lines.append(f'<outline text="步骤：{escape(cell_text(c.steps))}"/>')
                if c.expected:
                    lines.append(f'<outline text="预期：{escape(cell_text(c.expected))}"/>')
                if c.priority:
                    lines.append(f'<outline text="优先级：{c.priority.value}"/>')
                
                lines.append("</outline>")
            lines.append("</outline>")
        
        lines.extend(["</outline>", "</body>", "</opml>"])
        return "\n".join(lines).encode("utf-8")
