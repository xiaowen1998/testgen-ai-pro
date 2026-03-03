"""
AI 测试用例生成器 V7.0
测试点驱动双阶段 · 精简UI：测试点清单 → 用例表格 → 分页（无筛选/统计）

【耗时分布】主要耗时在：
1. 第1/2阶段 API 请求（网络+模型生成），占绝大部分时间；
2. post_process 中的 dedupe_by_title_similarity（已优化：长度预判+quick_ratio 减少 ratio 调用）；
3. parse_llm_response 仅在含前置说明或代码块时走括号匹配，通常直接 json.loads。
40 个测试点仅 1 批请求，无并行；若仍慢可检查网络或模型侧时延。
"""

import base64
import hashlib
import io
import os
import json
import re
import time
import difflib
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Optional, Any, Tuple

import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

st.set_page_config(
    page_title=" TestGen AI Testcases",
    page_icon="🧪",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ========== 常量 ==========
PER_PAGE_DEFAULT = 15
PER_PAGE_OPTIONS = [10, 15, 20, 30]  # 每页条数可选
TABLE_HEADERS = ["用例编号", "用例名称", "所属需求模块", "测试类型", "前置条件", "测试步骤", "测试数据", "预期结果", "优先级"]
# XMind 解析后用于生成的最大字符数，避免内容过多导致测试点爆炸
XMIND_MAX_EXTRACT_CHARS = 20000
# 第 2 阶段单次请求最多携带的测试点数量，超出则分批生成（每批 15 便于进度展示）
STAGE2_MAX_POINTS_PER_BATCH = 15
# 第 1 阶段需求文本最大长度（字符）
STAGE1_MAX_CONTENT_CHARS = 12000
# 单次 API 请求超时（秒）；生成用例时输出较长，默认 180。仍超时可设环境变量 API_REQUEST_TIMEOUT=300
API_REQUEST_TIMEOUT = 180
PRIORITY_ORDER = {"高": 0, "中": 1, "低": 2}

# 类型标签配置：多别名映射到 (显示文本, CSS类名, 颜色)
TYPE_TAG_CONFIG: dict[tuple[str, ...], tuple[str, str, str]] = {
    ("功能", "功能测试"): ("功能", "tag-func", "#3b82f6"),
    ("边界", "边界值", "边界测试"): ("边界", "tag-boundary", "#f97316"),
    ("异常", "异常测试"): ("异常", "tag-exception", "#ef4444"),
    ("兼容性", "兼容"): ("兼容性", "tag-compatibility", "#6b7280"),
    ("性能", "性能测试"): ("性能", "tag-performance", "#8b5cf6"),
    ("安全", "安全测试"): ("安全", "tag-security", "#374151"),
    ("UI", "UI测试", "界面", "界面测试"): ("UI", "tag-ui", "#06b6d4"),
    ("接口", "API", "接口测试", "API测试"): ("接口", "tag-api", "#10b981"),
    ("冒烟", "冒烟测试"): ("冒烟", "tag-smoke", "#eab308"),
    ("回归", "回归测试"): ("回归", "tag-regression", "#ec4899"),
}


def get_type_tag_style(test_type: str) -> tuple[str, str, str]:
    """
    根据测试类型获取显示标签、CSS 类和颜色。
    支持别名与模糊匹配（如「功能测试」→ 功能）。
    返回: (显示文本, CSS类名, 颜色值)
    """
    if not test_type or not str(test_type).strip():
        return ("其他", "tag-other", "#9ca3af")
    test_type_lower = str(test_type).strip().lower()
    for aliases, (display, css_class, color) in TYPE_TAG_CONFIG.items():
        for alias in aliases:
            if alias.lower() in test_type_lower or test_type_lower in alias.lower():
                return (display, css_class, color)
    return (str(test_type).strip(), "tag-other", "#9ca3af")


def _build_type_tag_class() -> dict[str, str]:
    """构建类型/别名的 CSS 类映射，供兼容旧用法。"""
    mapping: dict[str, str] = {}
    for aliases, (_display, css_class, _color) in TYPE_TAG_CONFIG.items():
        for alias in aliases:
            mapping[alias] = css_class
    return mapping


TYPE_TAG_CLASS = _build_type_tag_class()

# 编辑态类型下拉选项（与 TYPE_TAG_CONFIG 显示文本一致）
TYPE_OPTS = ["功能", "边界", "异常", "兼容性", "性能", "安全", "UI", "接口", "冒烟", "回归"]

# V6.0 第1阶段：测试点分析
PROMPT_STAGE1_TEST_POINTS = """你是一位资深测试架构师。请先深度分析以下需求文档，识别所有需要验证的测试点。

需求内容：
{content}

要求：
1. 逐条阅读需求，识别每个功能规则、状态变化、交互逻辑、数据计算、展示规则、异常分支
2. 输出完整的"测试点清单"，每个测试点一句话描述（如："验证温层标签在冷冻商品上的展示"）
3. 只输出测试点清单，不生成用例
4. 确保没有遗漏任何需求点，有多少列多少

输出格式（严格按此格式，便于解析）：
测试点1：XXXX
测试点2：XXXX
测试点3：XXXX
..."""

# V6.0 第2阶段：基于测试点生成用例（强制：用例数 ≥ 测试点数）
PROMPT_STAGE2_CASES = """基于以下测试点清单，为每个测试点生成对应的测试用例。

【测试点清单】
{test_points}

【绝对强制要求】
1. ⚠️ **每个测试点至少生成 1 条用例，复杂测试点可生成 2～3 条**（正常流程+边界+异常）
2. **用例总数必须 ≥ 测试点数量**（本批共 {N} 个测试点，至少生成 {N} 条用例）
3. 严禁合并、省略或跳过任何一个测试点
4. 复杂测试点（多状态/计算/异常）必须生成多条用例覆盖不同场景
5. test_point 字段必须填写对应序号（如 "测试点1"、"测试点2"）

【输出要求】
- 仅输出一个 JSON 数组，不要 markdown 代码块或其它说明
- 数组长度必须 ≥ {N}；若长度 < {N} 则不合格，需重新生成
- 建议生成数量：{N}～{N2} 条用例

【数量自检】
生成后自检：数组长度是否 ≥ {N}？test_point 是否覆盖测试点1～{N}？

格式示例：
[{{"case_id":"TC001","case_name":"标题","module":"购物车","test_point":"测试点1","test_type":"功能","precondition":"已登录","steps":"步骤1；步骤2","test_data":"具体数据","expected":"预期表现","priority":"高"}}]"""

SYSTEM_STAGE2 = """你生成测试用例时必须：
1. 每个测试点至少 1 条对应用例，用例总数 ≥ 测试点数量，禁止少生成
2. test_point 填写 "测试点N"（N 为序号），步骤具体到 UI 元素，数据给具体值
3. 结果可断言，前置条件完整；复杂测试点生成 2～3 条
只输出有效的 JSON 数组，不要其他说明。"""


def _extract_json_array_from_text(text: str) -> Optional[str]:
    """从可能含前置说明或 markdown 的文本中提取 JSON 数组字符串。"""
    if not text or not text.strip():
        return None
    text = text.strip()
    # 1) 去掉 markdown 代码块
    if "```" in text:
        for pattern in (r"```(?:json)?\s*([\s\S]*?)\s*```", r"```\s*([\s\S]*?)\s*```"):
            m = re.search(pattern, text)
            if m:
                raw = m.group(1).strip()
                if raw.startswith("["):
                    return raw
    # 2) 直接整段解析
    if text.startswith("["):
        return text
    # 3) 查找第一个 '[' 并匹配到对应的 ']'（仅双引号视为字符串边界，符合 JSON）
    start = text.find("[")
    if start == -1:
        return None
    depth = 0
    in_str = False
    escape = False
    for i in range(start, len(text)):
        c = text[i]
        if escape:
            escape = False
            continue
        if in_str:
            if c == "\\":
                escape = True
            elif c == '"':
                in_str = False
            continue
        if c == '"':
            in_str = True
        elif c == "[":
            depth += 1
        elif c == "]":
            depth -= 1
            if depth == 0:
                return text[start : i + 1]
    return None


def parse_llm_response(text: str) -> List[dict]:
    """解析模型返回的 JSON：支持代码块、尾部逗号、list/dict 与 test_cases 字段，无冗余括号扫描。"""
    if not text or not text.strip():
        return []
    text = text.strip()
    if text.startswith("```"):
        m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text)
        if m:
            text = m.group(1).strip()
    try:
        cleaned = text.replace(",]", "]").replace(",}", "}")
        data = json.loads(cleaned)
        if isinstance(data, list):
            return [c for c in data if isinstance(c, dict)]
        if isinstance(data, dict):
            if "test_cases" in data:
                return [c for c in data["test_cases"] if isinstance(c, dict)]
            return [data]
    except json.JSONDecodeError:
        pass
    return []


def normalize_case(c: dict) -> dict:
    """统一字段名：module/所属需求模块、test_point 等"""
    out = dict(c)
    if "所属需求模块" in out and "module" not in out:
        out["module"] = out.get("所属需求模块") or ""
    if "module" not in out:
        out["module"] = out.get("所属需求模块") or out.get("test_point") or "-"
    return out


def dedupe_by_title_similarity(cases: List[dict], threshold: float = 0.92) -> List[dict]:
    """去重：先哈希去重 O(n)，≥20 条再按长度分桶只比较相邻桶，降低 O(n²) 到近似 O(n)。"""
    if not cases or len(cases) < 2:
        return cases

    # 第一层：哈希去重 O(n)
    seen_names = set()
    unique_cases = []
    for c in cases:
        name = (c.get("case_name") or "").strip()
        if not name:
            unique_cases.append(c)
            continue
        if name in seen_names:
            continue
        seen_names.add(name)
        unique_cases.append(c)

    # <20 条跳过相似度检查
    if len(unique_cases) < 20:
        return unique_cases

    # 第二层：按长度分桶，只比较相邻桶
    buckets: dict[int, list[tuple[dict, str]]] = {}
    out = []

    for c in unique_cases:
        name = (c.get("case_name") or "").strip()
        if not name:
            out.append(c)
            continue

        bucket_key = len(name) // 10
        candidates: list[tuple[dict, str]] = []
        for bk in [bucket_key - 1, bucket_key, bucket_key + 1]:
            if bk in buckets:
                candidates.extend(buckets[bk])

        is_dup = False
        for _oc, oname in candidates:
            max_len = max(len(name), len(oname))
            if max_len == 0 or abs(len(name) - len(oname)) / max_len > 0.15:
                continue
            sm = difflib.SequenceMatcher(None, name, oname)
            if sm.quick_ratio() < threshold:
                continue
            if sm.ratio() >= threshold:
                is_dup = True
                break

        if not is_dup:
            if bucket_key not in buckets:
                buckets[bucket_key] = []
            buckets[bucket_key].append((c, name))
            out.append(c)

    return out


def sort_by_priority(cases: List[dict]) -> List[dict]:
    def key(c):
        return PRIORITY_ORDER.get((c.get("priority") or "中").strip(), 1)
    return sorted(cases, key=key)


def _track_coverage(cases: List[dict], test_points: List[str]) -> dict[int, bool]:
    """统计每个测试点是否至少被 1 条用例覆盖。返回 索引(0-based) -> 是否覆盖。"""
    coverage: dict[int, bool] = {i: False for i in range(len(test_points))}
    for c in cases:
        tp = str(c.get("test_point", ""))
        m = re.search(r"测试点\s*(\d+)", tp)
        if m:
            point_num = int(m.group(1))
            idx = point_num - 1
            if 0 <= idx < len(test_points):
                coverage[idx] = True
    return coverage


def create_placeholder_case(test_point: str, point_num: int) -> dict:
    """创建占位用例，供用户手动补充。"""
    return {
        "case_id": f"TC{point_num:03d}",
        "case_name": f"【待补充】测试点{point_num}",
        "module": "未分类",
        "test_point": f"测试点{point_num}",
        "test_type": "功能",
        "precondition": "【请手动补充】",
        "steps": f"【请手动补充】\n\n原始测试点：{(test_point or '')[:100]}",
        "test_data": "【请手动补充】",
        "expected": "【请手动补充】",
        "priority": "中",
        "_is_placeholder": True,
    }


def fill_missing_cases(
    client: OpenAI,
    test_points: List[str],
    content: str,
    cases: List[dict],
    coverage_map: dict[int, bool],
) -> List[dict]:
    """仅针对未覆盖的测试点单独请求生成，不重跑全量。"""
    missing_indices = [i for i in range(len(test_points)) if not coverage_map.get(i, False)]
    if not missing_indices:
        return cases
    st.info(f"🔄 检测到 {len(missing_indices)} 个测试点未覆盖，正在精准补全…")
    for idx in missing_indices:
        point = test_points[idx]
        point_num = idx + 1
        with st.spinner(f"补全测试点{point_num}：{point[:24]}…"):
            new_cases: List[dict] = []
            for attempt in range(3):
                try:
                    raw = _run_stage2_one_batch(client, [point], content, model_override=None)
                    if raw:
                        for c in raw:
                            c["test_point"] = f"测试点{point_num}"
                            new_cases.append(normalize_case(c))
                        break
                except Exception:
                    time.sleep(1)
            if new_cases:
                cases.append(new_cases[0])
            else:
                cases.append(create_placeholder_case(point, point_num))
    return cases


def strict_validate_cases(cases: List[dict], test_points: List[str]) -> Tuple[List[dict], bool]:
    """
    严格校验：用例数必须 ≥ 测试点数，且每个测试点至少被覆盖 1 次。
    返回: (用例列表, 是否通过)
    """
    case_count = len(cases)
    point_count = len(test_points)
    if case_count < point_count:
        missing = point_count - case_count
        st.error(f"❌ 用例数量({case_count}) < 测试点数量({point_count})，少了 {missing} 条。可能原因：模型漏生成、JSON 解析丢失或去重过度。")
        return cases, False
    covered_points = set()
    for c in cases:
        tp = str(c.get("test_point", ""))
        m = re.search(r"测试点\s*(\d+)", tp)
        if m:
            covered_points.add(int(m.group(1)))
    expected_points = set(range(1, point_count + 1))
    missing_points = expected_points - covered_points
    if missing_points:
        st.error(f"❌ 测试点 {sorted(missing_points)} 没有用例覆盖。")
        return cases, False
    if case_count == point_count:
        st.info(f"ℹ️ 生成 {case_count} 条用例，每个测试点 1 条（复杂场景建议生成多条以提升覆盖）。")
    else:
        st.success(f"✅ 生成 {case_count} 条用例，覆盖 {point_count} 个测试点，平均每测试点 {case_count / point_count:.1f} 条。")
    return cases, True


def dedupe_protect_test_points(cases: List[dict], test_points: List[str]) -> List[dict]:
    """去重时保护：按 test_point 分组，组内按 case_name 去重，确保每个测试点至少保留 1 条。"""
    original_cases = list(cases)
    groups: dict[str, list[dict]] = {}
    for c in cases:
        tp = str(c.get("test_point", "未知")).strip()
        if tp not in groups:
            groups[tp] = []
        groups[tp].append(c)
    result: list[dict] = []
    for _tp, group_cases in groups.items():
        if len(group_cases) == 1:
            result.append(group_cases[0])
            continue
        seen_names: set[str] = set()
        result.append(group_cases[0])
        name0 = (group_cases[0].get("case_name") or "").strip()
        seen_names.add(name0)
        for c in group_cases[1:]:
            name = (c.get("case_name") or "").strip()
            if name not in seen_names:
                seen_names.add(name)
                result.append(c)
    for i in range(1, len(test_points) + 1):
        tp_key = f"测试点{i}"
        if not any(str(c.get("test_point", "")).strip() == tp_key for c in result):
            for c in original_cases:
                if str(c.get("test_point", "")).strip() == tp_key:
                    result.append(c)
                    break
    return result


def post_process(cases: List[dict], test_points: Optional[List[str]] = None) -> List[dict]:
    """后处理；传入 test_points 时使用「按测试点保护」的去重，保证用例数 ≥ 测试点数。"""
    cases = [normalize_case(c) for c in cases]
    if test_points is not None:
        cases = dedupe_protect_test_points(cases, test_points)
    else:
        if len(cases) > 150:
            seen = set()
            unique = []
            for c in cases:
                name = (c.get("case_name") or "").strip()
                if name not in seen:
                    seen.add(name)
                    unique.append(c)
            cases = unique
        else:
            cases = dedupe_by_title_similarity(cases, 0.92)
    cases = sort_by_priority(cases)
    for i, c in enumerate(cases, 1):
        c["case_id"] = f"TC{i:03d}"
    return cases


def run_stage1_test_points(client: OpenAI, content: str) -> List[str]:
    """第1阶段：分析需求，返回测试点清单。支持多模型回退（额度用尽时自动换下一个）。"""
    content_limited = content[:STAGE1_MAX_CONTENT_CHARS]
    if len(content) > STAGE1_MAX_CONTENT_CHARS:
        content_limited += "\n\n（需求已截断，仅前 {} 字参与分析。）".format(STAGE1_MAX_CONTENT_CHARS)
    prompt = PROMPT_STAGE1_TEST_POINTS.format(content=content_limited)
    response, _ = _chat_create_with_model_fallback(
        client,
        messages=[
            {"role": "system", "content": "你只输出测试点清单，格式为 测试点N：描述，不要其他内容。"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
    )
    text = (response.choices[0].message.content or "").strip()
    points = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        # 匹配 "测试点1：xxx" 或 "1. xxx"
        m = re.match(r"测试点\s*\d+\s*[：:]\s*(.+)", line, re.I)
        if m:
            points.append(m.group(1).strip())
        elif re.match(r"\d+[\.．、]\s*.+", line):
            points.append(re.sub(r"^\d+[\.．、]\s*", "", line).strip())
    return points


def _run_stage2_one_batch(
    client: OpenAI,
    test_points_batch: List[str],
    _content: str,
    model_override: Optional[str] = None,
) -> List[dict]:
    """执行第2阶段的一批测试点。model_override 指定时用该模型单次请求，否则走模型列表 fallback。"""
    batch_text = "\n".join(f"测试点{i+1}：{p}" for i, p in enumerate(test_points_batch))
    n = len(test_points_batch)
    prompt = PROMPT_STAGE2_CASES.format(test_points=batch_text, N=n, N2=n * 2)
    messages = [
        {"role": "system", "content": SYSTEM_STAGE2},
        {"role": "user", "content": prompt},
    ]
    if model_override:
        response, _ = _chat_create_with_model(
            client, model_override, messages, temperature=0.4
        )
    else:
        response, _ = _chat_create_with_model_fallback(
            client, messages, temperature=0.4
        )
    text = (response.choices[0].message.content or "").strip()
    cases = parse_llm_response(text)
    if not isinstance(cases, list):
        cases = [cases] if isinstance(cases, dict) else []
    return cases


def run_stage2_cases(client: OpenAI, test_points: List[str], content: str) -> List[dict]:
    """第2阶段：≤15 条单次请求；多批时每批 15 条、最多 3 批并发，主线程更新进度条。"""
    if len(test_points) <= STAGE2_MAX_POINTS_PER_BATCH:
        all_cases = _run_stage2_one_batch(client, test_points, content, model_override=None)
        cases = post_process(all_cases, test_points)
        coverage_map = _track_coverage(cases, test_points)
        return cases, coverage_map

    batches = [
        test_points[i : i + STAGE2_MAX_POINTS_PER_BATCH]
        for i in range(0, len(test_points), STAGE2_MAX_POINTS_PER_BATCH)
    ]
    if not batches:
        return []

    progress_bar = st.progress(0.0)
    status_text = st.empty()
    model_list = get_qwen_model_list()
    max_workers = min(len(batches), 3)
    results_by_index = {}
    completed = 0

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_idx = {
            executor.submit(
                _run_stage2_one_batch,
                client,
                batch,
                content,
                model_override=model_list[i % len(model_list)],
            ): i
            for i, batch in enumerate(batches)
        }
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                results_by_index[idx] = future.result()
            except Exception:
                results_by_index[idx] = []
            completed += 1
            status_text.caption(f"已完成 {completed}/{len(batches)} 批（{completed * 100 // len(batches)}%）")
            progress_bar.progress(completed / len(batches))

    progress_bar.empty()
    status_text.empty()

    all_cases = []
    for i in range(len(batches)):
        all_cases.extend(results_by_index.get(i, []))
    cases = post_process(all_cases, test_points)
    coverage_map = _track_coverage(cases, test_points)
    return cases, coverage_map


# ========== 文档解析 ==========
def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = [p.extract_text() for p in reader.pages if p.extract_text()]
        return "\n\n".join(parts).strip() if parts else ""
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败: {e}") from e


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
        if not file_bytes or len(file_bytes) < 4:
            raise RuntimeError("文件为空或过短，无法解析。")
        if file_bytes[:2] != b"PK":
            raise RuntimeError(
                "该文件不是有效的 .docx 格式（.docx 应为 ZIP 结构）。"
                "请确认：1) 文件为 Office 2007 及以上另存为的 .docx；2) 若为旧版 .doc，请用 Word 另存为 .docx 后再上传。"
            )
        buf = io.BytesIO(file_bytes)
        buf.seek(0)
        doc = Document(buf)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"Word 解析失败: {e}") from e


def extract_txt_or_md(file_bytes: bytes, _filename: str = "") -> str:
    for enc in ("utf-8", "gbk", "gb2312"):
        try:
            return file_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    raise RuntimeError("无法识别文件编码。")


def _xmind_topic_to_text(node: dict, parts: List[str]) -> None:
    """递归提取 XMind 节点标题与子节点文本。"""
    if not isinstance(node, dict):
        return
    title = node.get("title") or node.get("text") or ""
    if isinstance(title, str) and title.strip():
        parts.append(title.strip())
    children = node.get("children")
    if isinstance(children, dict):
        for key in ("attached", "summary", "floating"):
            arr = children.get(key)
            if isinstance(arr, list):
                for c in arr:
                    _xmind_topic_to_text(c, parts)
    elif isinstance(children, list):
        for c in children:
            _xmind_topic_to_text(c, parts)


def extract_xmind_text(file_bytes: bytes) -> str:
    """从 .xmind 文件（ZIP 内 content.json）解析出纯文本。"""
    import zipfile
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as z:
            # 优先 content.json（XMind 8）
            names = z.namelist()
            content_name = None
            for n in names:
                if "content.json" in n or n == "content.json":
                    content_name = n
                    break
            if not content_name:
                for n in names:
                    if n.endswith(".json") and "content" in n.lower():
                        content_name = n
                        break
            if not content_name:
                raise RuntimeError("未在 XMind 文件中找到 content.json。")
            with z.open(content_name) as f:
                data = json.load(f)
        parts = []
        # 可能为单对象或数组（多 sheet）
        if isinstance(data, list):
            for sheet in data:
                root = sheet.get("rootTopic") if isinstance(sheet, dict) else None
                if root:
                    _xmind_topic_to_text(root, parts)
        elif isinstance(data, dict) and data.get("rootTopic"):
            _xmind_topic_to_text(data["rootTopic"], parts)
        text = "\n".join(parts).strip()
        if not text:
            raise RuntimeError("XMind 中未解析出有效文本。")
        # 限制长度，避免内容过多导致测试点过多、生成失败
        if len(text) > XMIND_MAX_EXTRACT_CHARS:
            text = text[:XMIND_MAX_EXTRACT_CHARS] + "\n\n（以上内容已截断，仅保留前 {} 字用于生成。如需完整需求请分段上传或精简脑图。）".format(XMIND_MAX_EXTRACT_CHARS)
        return text
    except zipfile.BadZipFile as e:
        raise RuntimeError(f"不是有效的 XMind/ZIP 文件: {e}") from e
    except json.JSONDecodeError as e:
        raise RuntimeError(f"XMind content.json 解析失败: {e}") from e


def extract_zip_text(file_bytes: bytes) -> str:
    """解析 ZIP 压缩包内支持的文档（PDF/Word/TXT/MD/XMind），汇总文本。"""
    import zipfile
    out_parts = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as z:
            for name in z.namelist():
                if name.startswith("__MACOSX") or "/." in name or name.endswith("/"):
                    continue
                base = (name.split("/")[-1] or name).lower()
                try:
                    raw = z.read(name)
                except Exception:
                    continue
                if not raw or len(raw) < 2:
                    continue
                try:
                    if base.endswith(".pdf"):
                        out_parts.append(extract_pdf_text(raw))
                    elif base.endswith(".docx"):
                        out_parts.append(extract_docx_text(raw))
                    elif base.endswith((".txt", ".md", ".markdown")):
                        out_parts.append(extract_txt_or_md(raw, base))
                    elif base.endswith(".xmind"):
                        out_parts.append(extract_xmind_text(raw))
                    else:
                        continue
                except Exception:
                    continue
        text = "\n\n".join(p for p in out_parts if (p or "").strip()).strip()
        if not text:
            raise RuntimeError("ZIP 内未找到可解析的 PDF/Word/TXT/MD/XMind 文件。")
        return text
    except zipfile.BadZipFile as e:
        raise RuntimeError(f"不是有效的 ZIP 文件: {e}") from e


def extract_text_from_upload(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    if name.endswith(".pdf"):
        return extract_pdf_text(raw)
    if name.endswith(".docx"):
        return extract_docx_text(raw)
    if name.endswith(".txt") or name.endswith(".md") or name.endswith(".markdown"):
        return extract_txt_or_md(raw, name)
    if name.endswith(".xmind"):
        return extract_xmind_text(raw)
    if name.endswith(".zip"):
        return extract_zip_text(raw)
    raise ValueError("仅支持 PDF、Word(.docx)、TXT、Markdown、XMind(.xmind)、ZIP 压缩包。")


# ========== 图片识别（多模态，根据截图自动识别类型并提取）==========
IMAGE_PROMPT_UNIVERSAL = """请根据图片内容自动判断类型（文本文档、UI 设计图、流程图、手写笔记等），并提取其中所有可用于需求分析的信息：
- 若是文档/截图：识别全部文字，保持格式与结构（标题、段落、列表、表格）。
- 若是 UI 设计图：提取页面模块、控件与文案、交互说明、业务规则与异常提示。
- 若是流程图：提取起止节点、步骤、分支与异常分支。
- 若是手写：尽量辨认字迹并保持段落结构，不确定处用 [?] 标记。
只输出提取后的结构化文本，不要解释图片类型或添加多余说明。"""


def extract_text_from_image(uploaded_file) -> str:
    """根据截图自动识别类型并提取文本与需求（多模态模型）。"""
    image_bytes = uploaded_file.getvalue()
    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    file_type = getattr(uploaded_file, "type", None) or "image/jpeg"
    mime_type = file_type if str(file_type).startswith("image/") else "image/jpeg"

    client = get_qwen_client()
    if not client:
        raise RuntimeError("未配置 QWEN_API_KEY")

    prompt = IMAGE_PROMPT_UNIVERSAL

    messages = [
        {"role": "system", "content": "你是一个专业的需求文档识别助手，擅长从图片中提取结构化信息。"},
        {
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}},
                {"type": "text", "text": prompt},
            ],
        },
    ]

    timeout = int(os.getenv("API_REQUEST_TIMEOUT", "60"))
    try:
        response = client.chat.completions.create(
            model="qwen-vl-max",
            messages=messages,
            temperature=0.3,
            max_tokens=2000,
            timeout=timeout,
        )
        text = (response.choices[0].message.content or "").strip()
        return text
    except Exception as e:
        err = str(e).lower()
        if "timeout" in err:
            raise RuntimeError("图片识别超时，请尝试压缩图片或稍后重试。") from e
        if "rate limit" in err or "429" in err:
            raise RuntimeError("视觉模型调用频繁，请稍后重试。") from e
        raise RuntimeError(f"图片识别失败：{e}") from e


# 抓取到疑似“未登录”页面时的提示（引导用户复制文本或上传文档）
MSG_LOGIN_REQUIRED = (
    "页面未登录，无法通过 URL 自动登录。请复制页面正文到「文本输入」中生成，或使用「文档上传」上传文档。"
)
# 抓取结果为“需启用 JavaScript”等无效内容时的提示
MSG_JS_NO_CONTENT = (
    "该页面需在浏览器中执行 JavaScript 后才显示正文，当前无法抓取到有效内容。请将页面正文复制到「文本输入」中，或使用「文档上传」上传文档。"
)


def _is_login_required_page(text: str) -> bool:
    """根据抓取结果判断是否为需登录页面（内容过短且含登录相关关键词）。"""
    if not text or len(text) > 800:
        return False
    t = text.strip().lower()
    keywords = ("登录", "请登录", "未登录", "login", "sign in", "登录以继续", "去登录", "立即登录")
    return any(kw.lower() in t or kw in text for kw in keywords)


def _is_js_placeholder_content(text: str) -> bool:
    """抓取结果是否为“需启用 JavaScript”等无效占位内容（需给出提示而非静默填入）。"""
    if not text or len(text) > 600:
        return False
    t = text.strip().lower()
    return "enable javascript" in t or "you need to enable javascript" in t or "请启用 javascript" in t


def _extract_text_from_html(html: str) -> str:
    """从 HTML 中提取正文，去掉 script/style 等。"""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    body = soup.find("body") or soup
    text = body.get_text(separator="\n", strip=True) if body else ""
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _fetch_url_with_playwright(url: str) -> str:
    """使用 Playwright 抓取需要 JavaScript 渲染的页面。"""
    from playwright.sync_api import sync_playwright
    url = url.strip()
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            page = browser.new_page()
            # 使用 load 而非 networkidle，避免 SPA/长轮询页面一直不满足 networkidle 导致超时
            page.goto(url, wait_until="load", timeout=25000)
            page.wait_for_timeout(4000)
            html = page.content()
            return _extract_text_from_html(html)
        finally:
            browser.close()


def fetch_url_text(url: str) -> str:
    """抓取网页正文；若为 JS 渲染页面则自动尝试 Playwright。"""
    try:
        import requests
        r = requests.get(
            url.strip(),
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"},
            timeout=15,
        )
        r.raise_for_status()
        r.encoding = r.apparent_encoding or "utf-8"
        text = _extract_text_from_html(r.text)
        if not text or len(text) < 80 or "enable javascript" in text.lower() or "enable JavaScript" in text:
            try:
                text = _fetch_url_with_playwright(url)
            except ImportError:
                raise RuntimeError(
                    "该页面需执行 JavaScript 才能显示内容。请安装 Playwright 后重试："
                    " pip install playwright && playwright install chromium"
                    "；或直接将页面正文复制到「文本输入」中。"
                ) from None
            except Exception as e:
                err_msg = str(e).strip()
                if "timeout" in err_msg.lower() or "Timeout" in err_msg:
                    raise RuntimeError(
                        "页面加载超时（该链接可能需登录或加载较慢）。请将正文复制到「文本输入」中生成，或使用「文档上传」上传文档。"
                    ) from e
                raise RuntimeError(f"Playwright 抓取失败: {e}") from e
        if _is_login_required_page(text):
            raise RuntimeError(MSG_LOGIN_REQUIRED)
        if _is_js_placeholder_content(text):
            raise RuntimeError(MSG_JS_NO_CONTENT)
        return text
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"抓取失败: {e}") from e


def get_qwen_client() -> Optional[OpenAI]:
    api_key = (os.getenv("QWEN_API_KEY") or "").strip()
    if not api_key:
        return None
    return OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
        timeout=float(os.getenv("API_REQUEST_TIMEOUT", API_REQUEST_TIMEOUT)),
    )


def get_qwen_model_list() -> List[str]:
    """从环境变量 QWEN_MODEL_LIST 读取模型列表（逗号分隔），第一个用完后抛异常则尝试下一个。默认 qwen-turbo,qwen-plus,qwen-max。"""
    raw = (os.getenv("QWEN_MODEL_LIST") or os.getenv("QWEN_MODEL") or "qwen-turbo").strip()
    if "," in raw:
        return [m.strip() for m in raw.split(",") if m.strip()]
    return [raw] if raw else ["qwen-turbo"]


# 智能模型路由：按输入类型选择展示与说明（图片识别用 qwen-vl-max，后续阶段用文本模型）
MODEL_ROUTING = {
    "text": ("qwen-turbo", "qwen-turbo", "⚡ 极速模式", "文本分析，成本低速度快"),
    "document": ("qwen-turbo", "qwen-turbo", "⚡ 极速模式", "文档解析已完成"),
    "url": ("qwen-turbo", "qwen-turbo", "⚡ 极速模式", "网页文本分析"),
    "image": ("qwen-vl-max", "qwen-turbo", "🖼️ 视觉模式", "图片已识别为文本，将用文本模型生成测试点与用例"),
}


def get_models_for_input(input_source: str) -> Tuple[str, str, str, str]:
    """根据输入类型返回 (stage1 模型, stage2 模型, 显示名称, 描述)。"""
    return MODEL_ROUTING.get(input_source, ("qwen-turbo", "qwen-turbo", "⚡ 极速模式", ""))


def _is_quota_or_retryable_error(exc: Exception) -> bool:
    """是否为额度用尽/限流等可换模型重试的错误"""
    msg = str(exc).lower()
    if "429" in msg:
        return True
    for k in ("quota", "rate limit", "insufficient", "limit exceeded", "免费", "额度", "throttl"):
        if k in msg:
            return True
    return False


def _chat_create_with_model(
    client: OpenAI,
    model: str,
    messages: list,
    temperature: float = 0.3,
    **kwargs: Any,
) -> Tuple[Any, str]:
    """使用指定模型发起单次请求，不自动换模型。用于分批时每批绑定不同模型。"""
    resp = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        **kwargs,
    )
    return resp, model


def _chat_create_with_model_fallback(
    client: OpenAI,
    messages: list,
    temperature: float = 0.3,
    **kwargs: Any,
) -> Tuple[Any, str]:
    """
    按 QWEN_MODEL_LIST 顺序调用模型，当前模型报额度/限流类错误时自动尝试下一个，返回 (response, used_model)。
    """
    model_list = get_qwen_model_list()
    last_err: Optional[Exception] = None
    for model in model_list:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                **kwargs,
            )
            return resp, model
        except Exception as e:
            last_err = e
            if _is_api_key_error(e):
                raise
            if _is_quota_or_retryable_error(e):
                continue
            raise
    if last_err is not None:
        raise last_err
    raise RuntimeError("未配置可用模型列表（QWEN_MODEL_LIST）")


def _is_api_key_error(exc: Exception) -> bool:
    """判断是否为 API 密钥无效（401 / invalid_api_key）"""
    msg = str(exc).lower()
    return "401" in msg or "invalid_api_key" in msg or "incorrect api key" in msg


# ========== 表格与导出 ==========
def _cell_text(s: str) -> str:
    if not s:
        return "-"
    return str(s).replace("\n", " ").replace("|", "｜").strip() or "-"


def _get_module(c: dict) -> str:
    return _cell_text(c.get("module") or c.get("所属需求模块"))


def cases_to_rows(cases: List[dict]) -> List[List[str]]:
    return [
        [
            _cell_text(c.get("case_id")),
            _cell_text(c.get("case_name")),
            _get_module(c),
            _cell_text(c.get("test_type")),
            _cell_text(c.get("precondition")),
            _cell_text(c.get("steps")),
            _cell_text(c.get("test_data")),
            _cell_text(c.get("expected")),
            _cell_text(c.get("priority")),
        ]
        for c in cases
    ]


def to_markdown(cases: List[dict]) -> str:
    if not cases:
        return ""
    header_line = "| " + " | ".join(TABLE_HEADERS) + " |"
    sep = "| " + " | ".join(["---"] * len(TABLE_HEADERS)) + " |"
    rows = cases_to_rows(cases)
    body = "\n".join("| " + " | ".join(cell for cell in row) + " |" for row in rows)
    return "\n".join([header_line, sep, body])


def _html_esc(s: str) -> str:
    if not s:
        return ""
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def to_html_table_with_colors(
    cases: List[dict], start_idx: int = 0, include_thead: bool = True
) -> str:
    """一次性渲染用例表格为 HTML（10 列，无操作列），用于减少 Streamlit 组件数。"""
    if not cases:
        return ""
    headers = ["序号", "编号", "用例名称", "模块", "类型", "前置条件", "步骤", "测试数据", "预期结果", "优先级"]
    sb = ['<div class="table-scroll-wrap"><table class="cases-table">']
    if include_thead:
        sb.append("<thead><tr>")
        for h in headers:
            sb.append(f"<th>{h}</th>")
        sb.append("</tr></thead>")
    sb.append("<tbody>")
    for i, c in enumerate(cases):
        idx = start_idx + i + 1
        cid = str(c.get("case_id") or "-")
        name = str(c.get("case_name") or "-").replace("&", "&amp;").replace("<", "&lt;")
        module = str(c.get("module") or c.get("所属需求模块") or "-")
        t = str(c.get("test_type") or "-")
        display_text, tag_cls, _ = get_type_tag_style(t)
        t_esc = display_text.replace("&", "&amp;").replace("<", "&lt;")
        pre = str(c.get("precondition") or "-").replace("&", "&amp;").replace("<", "&lt;")
        steps = str(c.get("steps") or "-").replace("&", "&amp;").replace("<", "&lt;")
        data = str(c.get("test_data") or "-").replace("&", "&amp;").replace("<", "&lt;")
        exp = str(c.get("expected") or "-").replace("&", "&amp;").replace("<", "&lt;")
        p = str(c.get("priority") or "中")
        pri_cls = "tag-pri-high" if p == "高" else ("tag-pri-mid" if p == "中" else "tag-pri-low")
        row_cls = ' class="row-placeholder"' if c.get("_is_placeholder") else ""
        sb.append(f"<tr{row_cls}><td>{idx}</td><td>{cid}</td>")
        sb.append(f"<td class='cell-wrap'>{name}</td>")
        sb.append(f"<td>{module}</td>")
        sb.append(f'<td><span class="type-tag {tag_cls}">{t_esc}</span></td>')
        sb.append(f"<td class='cell-wrap'>{pre}</td>")
        sb.append(f"<td class='cell-wrap'>{steps}</td>")
        sb.append(f"<td class='cell-wrap'>{data}</td>")
        sb.append(f"<td class='cell-wrap'>{exp}</td>")
        sb.append(f'<td><span class="priority-tag {pri_cls}">{p}</span></td>')
        sb.append("</tr>")
    sb.append("</tbody></table></div>")
    return "".join(sb)


def to_html_table_with_type_colors(cases: List[dict], low_quality_ids: Optional[set] = None) -> str:
    low_quality_ids = low_quality_ids or set()
    if not cases:
        return ""
    sb = ['<table class="cases-table"><thead><tr><th>序号</th>']
    for h in TABLE_HEADERS:
        sb.append(f"<th>{_html_esc(h)}</th>")
    sb.append("</tr></thead><tbody>")
    for idx, c in enumerate(cases, 1):
        cid = _cell_text(c.get("case_id"))
        row_cls = ' class="row-low-quality"' if cid in low_quality_ids else ""
        sb.append(f"<tr{row_cls}>")
        sb.append(f"<td>{idx}</td>")
        sb.append(f"<td>{_html_esc(cid)}</td>")
        sb.append(f"<td>{_html_esc(_cell_text(c.get('case_name')))}</td>")
        sb.append(f"<td>{_html_esc(_get_module(c))}</td>")
        t = _cell_text(c.get("test_type"))
        display_text, cls, _ = get_type_tag_style(t)
        sb.append(f'<td><span class="type-tag {cls}">{_html_esc(display_text)}</span></td>')
        for key in ("precondition", "steps", "test_data", "expected", "priority"):
            sb.append(f"<td class='cell-wrap'>{_html_esc(_cell_text(c.get(key)))}</td>")
        sb.append("</tr>")
    sb.append("</tbody></table>")
    return "".join(sb)


def to_html_table_v7(cases: List[dict]) -> str:
    """V7: 表格含 操作 列，类型/优先级带颜色标签，支持横向滚动外层."""
    if not cases:
        return ""
    headers = ["序号", "编号", "用例名称", "模块", "类型", "前置条件", "步骤", "测试数据", "预期结果", "优先级", "操作"]
    sb = ['<div class="table-scroll-wrap"><table class="cases-table"><thead><tr>']
    for h in headers:
        sb.append(f"<th>{_html_esc(h)}</th>")
    sb.append("</tr></thead><tbody>")
    for idx, c in enumerate(cases, 1):
        cid = _cell_text(c.get("case_id"))
        sb.append("<tr>")
        sb.append(f"<td>{idx}</td>")
        sb.append(f"<td>{_html_esc(cid)}</td>")
        sb.append(f"<td class='cell-wrap'>{_html_esc(_cell_text(c.get('case_name')))}</td>")
        sb.append(f"<td>{_html_esc(_get_module(c))}</td>")
        t = _cell_text(c.get("test_type"))
        display_text, tag_cls, _ = get_type_tag_style(t)
        sb.append(f'<td><span class="type-tag {tag_cls}">{_html_esc(display_text)}</span></td>')
        for key in ("precondition", "steps", "test_data", "expected"):
            sb.append(f"<td class='cell-wrap'>{_html_esc(_cell_text(c.get(key)))}</td>")
        p = _cell_text(c.get("priority"))
        pri_cls = "tag-pri-high" if p == "高" else ("tag-pri-mid" if p == "中" else "tag-pri-low")
        sb.append(f'<td><span class="priority-tag {pri_cls}">{_html_esc(p)}</span></td>')
        sb.append("<td>编辑</td>")
        sb.append("</tr>")
    sb.append("</tbody></table></div>")
    return "".join(sb)


def to_excel_bytes(cases: List[dict]) -> bytes:
    import pandas as pd
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill
    df = pd.DataFrame(cases_to_rows(cases), columns=TABLE_HEADERS)
    buf = io.BytesIO()
    df.to_excel(buf, index=False, engine="openpyxl")
    buf.seek(0)
    wb = load_workbook(buf)
    ws = wb.active
    fill_high = PatternFill(start_color="FFCCCB", end_color="FFCCCB", fill_type="solid")
    fill_mid = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
    fill_low = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    pri_col = 9
    for row in range(2, ws.max_row + 1):
        val = (ws.cell(row=row, column=pri_col).value or "").strip()
        if val == "高":
            for col in range(1, 10):
                ws.cell(row=row, column=col).fill = fill_high
        elif val == "中":
            for col in range(1, 10):
                ws.cell(row=row, column=col).fill = fill_mid
        elif val == "低":
            for col in range(1, 10):
                ws.cell(row=row, column=col).fill = fill_low
    # 增加筛选：表头行启用自动筛选
    if ws.max_row >= 1:
        ws.auto_filter.ref = f"A1:I{ws.max_row}"
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out.read()


def to_word_bytes(cases: List[dict]) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("测试用例", 0)
    doc.add_paragraph("以下为根据需求生成的测试用例列表，可直接用于执行。")
    doc.add_paragraph()
    doc.add_heading("目录", level=1)
    for c in cases:
        doc.add_paragraph(f"{_cell_text(c.get('case_id'))} - {_cell_text(c.get('case_name'))}", style="List Bullet")
    doc.add_page_break()
    for i, c in enumerate(cases):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(f"{_cell_text(c.get('case_id'))} - {_cell_text(c.get('case_name'))}", level=1)
        doc.add_paragraph(f"所属模块：{_get_module(c)}")
        doc.add_paragraph(f"测试类型：{_cell_text(c.get('test_type'))} | 优先级：{_cell_text(c.get('priority'))}")
        doc.add_paragraph(f"前置条件：{_cell_text(c.get('precondition'))}")
        doc.add_paragraph(f"测试步骤：{_cell_text(c.get('steps'))}")
        doc.add_paragraph(f"测试数据：{_cell_text(c.get('test_data'))}")
        doc.add_paragraph(f"预期结果：{_cell_text(c.get('expected'))}")
    for s in doc.sections:
        s.header.is_linked_to_previous = False
        s.footer.is_linked_to_previous = False
        if not s.header.paragraphs:
            s.header.add_paragraph("TestGen AI V5.0 - 测试用例")
        else:
            s.header.paragraphs[0].text = "TestGen AI V5.0 - 测试用例"
        if not s.footer.paragraphs:
            s.footer.add_paragraph("测试用例文档")
        else:
            s.footer.paragraphs[0].text = "测试用例文档"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _find_chinese_font_path() -> Optional[str]:
    """返回可用于 PDF 中文的 TTF 字体路径（环境变量或系统常见路径）。"""
    env_path = (os.getenv("PDF_CHINESE_FONT") or os.getenv("REPORTLAB_CHINESE_FONT") or "").strip()
    if env_path and os.path.isfile(env_path):
        return env_path
    candidates = [
        os.path.expanduser("~/Library/Fonts/PingFang.ttc"),
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/msyhbd.ttf",
        "C:/Windows/Fonts/simsun.ttf",
        "C:/Windows/Fonts/simhei.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for p in candidates:
        if p and os.path.isfile(p):
            return p
    return None


def to_pdf_bytes(cases: List[dict]) -> bytes:
    """导出 PDF，优先使用中文字体以正确显示中文。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    font_name = "Helvetica"
    font_path = _find_chinese_font_path()
    if font_path:
        try:
            font_name = "PDFChinese"
            pdfmetrics.registerFont(TTFont(font_name, font_path))
        except Exception:
            font_name = "Helvetica"

    headers_short = ["编号", "名称", "模块", "类型", "前置", "步骤", "数据", "预期", "优先级"]
    col_widths = [22 * mm, 28 * mm, 22 * mm, 18 * mm, 22 * mm, 35 * mm, 22 * mm, 35 * mm, 18 * mm]
    data = [headers_short]
    for c in cases:
        row = [
            _cell_text(c.get("case_id"))[:12],
            _cell_text(c.get("case_name"))[:20],
            _get_module(c)[:12],
            _cell_text(c.get("test_type"))[:8],
            _cell_text(c.get("precondition"))[:14],
            _cell_text(c.get("steps"))[:24],
            _cell_text(c.get("test_data"))[:12],
            _cell_text(c.get("expected"))[:24],
            _cell_text(c.get("priority"))[:4],
        ]
        data.append(row)

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#667eea")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ]
        )
    )
    doc.build([t])
    buf.seek(0)
    return buf.read()


def _xmind_topic_id(prefix: str, n: int) -> str:
    """生成 XMind 节点 id，便于软件识别结构。"""
    return f"{prefix}{n}"


def to_xmind_bytes(cases: List[dict]) -> bytes:
    """将测试用例导出为 XMind 8/Zen 兼容格式（ZIP 内 content.json + manifest.json + metadata.json）。"""
    import zipfile

    by_module: dict = {}
    for c in cases:
        mod = _cell_text(c.get("module") or c.get("所属需求模块")) or "未分类"
        if mod not in by_module:
            by_module[mod] = []
        by_module[mod].append(c)

    node_counter = [0]

    def make_topic(title: str, children_attached: Optional[List[dict]] = None) -> dict:
        node_counter[0] += 1
        tid = _xmind_topic_id("topic", node_counter[0])
        node = {"id": tid, "title": title[:2000]}
        if children_attached:
            node["children"] = {"attached": children_attached}
        return node

    case_topics = []
    for mod, list_cases in by_module.items():
        mod_children = []
        for c in list_cases:
            cid = _cell_text(c.get("case_id"))
            name = _cell_text(c.get("case_name"))
            pre = _cell_text(c.get("precondition"))
            steps = _cell_text(c.get("steps"))
            data = _cell_text(c.get("test_data"))
            expected = _cell_text(c.get("expected"))
            pri = _cell_text(c.get("priority"))
            sub = []
            if pre and pre != "-":
                sub.append(make_topic(f"前置：{pre}"))
            if steps and steps != "-":
                sub.append(make_topic(f"步骤：{steps}"))
            if data and data != "-":
                sub.append(make_topic(f"数据：{data}"))
            if expected and expected != "-":
                sub.append(make_topic(f"预期：{expected}"))
            if pri and pri != "-":
                sub.append(make_topic(f"优先级：{pri}"))
            mod_children.append(make_topic(f"{cid} - {name}", sub if sub else None))
        case_topics.append(make_topic(mod, mod_children))

    node_counter[0] += 1
    root_id = _xmind_topic_id("root", node_counter[0])
    root = {"id": root_id, "title": "测试用例", "children": {"attached": case_topics}}
    # XMind 8/Zen：content.json 为数组，每项为 sheet，含 rootTopic
    sheet = {"id": "sheet-1", "title": "测试用例", "rootTopic": root}
    content = [sheet]
    content_bytes = json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")

    now_ts = int(time.time() * 1000)
    metadata = {
        "creator": {"name": "TestGen AI", "version": "7.0"},
        "created": {"timestamp": now_ts},
        "modified": {"timestamp": now_ts},
    }
    metadata_bytes = json.dumps(metadata, ensure_ascii=False).encode("utf-8")

    # manifest.json：含 checksum 可减少“文件可能已损坏”的误报（XMind 校验完整性）
    def _checksum(b: bytes) -> str:
        return hashlib.sha256(b).hexdigest()

    manifest = {
        "file-entries": {
            "content.json": {"checksum": _checksum(content_bytes)},
            "metadata.json": {"checksum": _checksum(metadata_bytes)},
        }
    }

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
        z.writestr("content.json", content_bytes.decode("utf-8"))
        z.writestr("metadata.json", metadata_bytes.decode("utf-8"))
    buf.seek(0)
    return buf.read()


def to_opml_bytes(cases: List[dict]) -> bytes:
    """导出为 OPML 格式（XMind、MindManager、幕布等均支持）。"""

    def escape_xml(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        "<opml version=\"2.0\">",
        "<head><title>测试用例</title></head>",
        "<body>",
        '<outline text="测试用例">',
    ]
    by_module: dict = {}
    for c in cases:
        mod = _cell_text(c.get("module") or c.get("所属需求模块")) or "未分类"
        if mod not in by_module:
            by_module[mod] = []
        by_module[mod].append(c)
    for mod, list_cases in by_module.items():
        lines.append(f'<outline text="{escape_xml(mod)}">')
        for c in list_cases:
            cid = _cell_text(c.get("case_id"))
            name = _cell_text(c.get("case_name"))
            title = f"{cid} - {name}"
            lines.append(f'<outline text="{escape_xml(title)}">')
            pre = _cell_text(c.get("precondition"))
            steps = _cell_text(c.get("steps"))
            exp = _cell_text(c.get("expected"))
            pri = _cell_text(c.get("priority"))
            if pre and pre != "-":
                lines.append(f'<outline text="前置：{escape_xml(pre)}"/>')
            if steps and steps != "-":
                lines.append(f'<outline text="步骤：{escape_xml(steps)}"/>')
            if exp and exp != "-":
                lines.append(f'<outline text="预期：{escape_xml(exp)}"/>')
            if pri and pri != "-":
                lines.append(f'<outline text="优先级：{escape_xml(pri)}"/>')
            lines.append("</outline>")
        lines.append("</outline>")
    lines.extend(["</outline>", "</body>", "</opml>"])
    return "\n".join(lines).encode("utf-8")


# ========== AI 评审 ==========
def run_ai_review(client: OpenAI, cases: List[dict], requirement: str) -> dict:
    """返回 {score, suggestions, low_quality_ids}"""
    cases_summary = json.dumps(
        [{"case_id": c.get("case_id"), "case_name": c.get("case_name"), "test_type": c.get("test_type"), "steps": c.get("steps"), "test_data": c.get("test_data"), "expected": c.get("expected")} for c in cases[:50]],
        ensure_ascii=False,
        indent=2,
    )
    prompt = f"""请对以下测试用例集合进行质量评审。

【需求摘要】
{requirement[:1500]}

【用例摘要（部分）】
{cases_summary}

【评审维度】
1. 覆盖度：是否遗漏重要需求点
2. 可执行性：步骤是否具体到可操作
3. 数据完整性：测试数据是否给出具体值
4. 预期可验证性：结果是否可判断通过/失败

请严格按以下 JSON 格式输出，不要其他内容：
{{
  "overall_score": 7,
  "coverage_score": 7,
  "executability_score": 8,
  "data_completeness_score": 6,
  "verifiability_score": 7,
  "suggestions": "优化建议文本，分条列出",
  "low_quality_case_ids": ["TC003", "TC007"]
}}
overall_score 为 1-10 分；low_quality_case_ids 为评分低于 6 分的用例编号列表。"""
    try:
        response, _ = _chat_create_with_model_fallback(
            client,
            messages=[
                {"role": "system", "content": "你只输出有效的 JSON 对象，不要 markdown 代码块。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )
        text = (response.choices[0].message.content or "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        data = json.loads(text)
        low = data.get("low_quality_case_ids") or []
        return {
            "overall_score": data.get("overall_score", 0),
            "suggestions": data.get("suggestions", ""),
            "low_quality_ids": set(str(x) for x in low),
        }
    except Exception:
        return {"overall_score": 0, "suggestions": "评审解析失败", "low_quality_ids": set()}


# ========== CSS ==========
CUSTOM_CSS = """
<style>
    .gradient-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #5b21b6 100%);
        padding: 1.5rem 2rem;
        border-radius: 16px;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 20px rgba(102, 126, 234, 0.35);
    }
    .gradient-header h1 { color: white !important; font-size: 2rem !important; font-weight: 700 !important; margin: 0 !important; }
    .gradient-header p { color: rgba(255,255,255,0.9) !important; font-size: 0.95rem !important; margin: 0.4rem 0 0 0 !important; }
    .stApp { background: linear-gradient(180deg, #f5f3ff 0%, #f8fafc 100%); }
    .pagination-info { font-size: 0.9rem; color: #64748b; margin-bottom: 0.5rem; }
    table.cases-table { border-collapse: collapse; width: 100%; border-radius: 10px; overflow: hidden; box-shadow: 0 2px 8px rgba(0,0,0,0.08); table-layout: fixed; }
    table.cases-table th, table.cases-table td { padding: 10px 12px; border-bottom: 1px solid #e5e7eb; word-wrap: break-word; }
    table.cases-table th { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important; color: white !important; font-weight: 600; }
    table.cases-table tbody tr:hover td { background: #f8fafc; }
    table.cases-table td.cell-wrap { white-space: pre-wrap; word-break: break-word; max-width: 180px; }
    table.cases-table tr.row-low-quality { background: #fef2f2 !important; }
    table.cases-table tr.row-placeholder { background: #fef3c7 !important; }
    .row-editing-wrap { background: #fef3c7 !important; padding: 8px 12px; border-radius: 8px; margin: 4px 0; border: 1px solid #f59e0b; }
    .type-tag { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 0.85em; font-weight: 500; }
    .type-tag.tag-func { background: #3b82f6; color: #fff; }
    .type-tag.tag-boundary { background: #f97316; color: #fff; }
    .type-tag.tag-exception { background: #ef4444; color: #fff; }
    .type-tag.tag-other { background: #9ca3af; color: #fff; }
    .type-tag.tag-compatibility { background: #6b7280; color: #fff; }
    .type-tag.tag-performance { background: #8b5cf6; color: #fff; }
    .type-tag.tag-security { background: #374151; color: #fff; }
    .type-tag.tag-ui { background: #06b6d4; color: #fff; }
    .type-tag.tag-api { background: #10b981; color: #fff; }
    .type-tag.tag-smoke { background: #eab308; color: #000; }
    .type-tag.tag-regression { background: #ec4899; color: #fff; }
    .priority-tag { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 0.85em; font-weight: 500; }
    .priority-tag.tag-pri-high { background: #fecaca; color: #991b1b; }
    .priority-tag.tag-pri-mid { background: #fef08a; color: #854d0e; }
    .priority-tag.tag-pri-low { background: #bbf7d0; color: #166534; }
    .table-scroll-wrap { overflow-x: auto; margin: 0.5rem 0; }
    .pagination-bar { display: flex; justify-content: flex-end; align-items: center; gap: 0.5rem; flex-wrap: wrap; font-size: 0.875rem; }
    .pagination-bar .stNumberInput { width: 3rem !important; }
    /* 底部分页紧凑：锚点在首列内，用 :has() 选中分页各列再缩小按钮/输入 */
    [data-testid="column"]:has(#pagination-compact),
    [data-testid="column"]:has(#pagination-compact) ~ [data-testid="column"] {
        padding-left: 0.15rem !important; padding-right: 0.15rem !important;
    }
    [data-testid="column"]:has(#pagination-compact) .stButton > button,
    [data-testid="column"]:has(#pagination-compact) ~ [data-testid="column"] .stButton > button {
        padding: 0.1rem 0.22rem !important; font-size: 0.7rem !important;
        min-height: 1.25rem !important; min-width: 1.4rem !important;
    }
    [data-testid="column"]:has(#pagination-compact) .stNumberInput input,
    [data-testid="column"]:has(#pagination-compact) ~ [data-testid="column"] .stNumberInput input {
        padding: 0.1rem 0.15rem !important; font-size: 0.7rem !important;
        width: 1.8rem !important; min-height: 1.25rem !important;
    }
    [data-testid="column"]:has(#pagination-compact) .stSelectbox > div,
    [data-testid="column"]:has(#pagination-compact) ~ [data-testid="column"] .stSelectbox > div {
        min-width: 2.8rem !important; font-size: 0.7rem !important;
    }
    [data-testid="column"]:has(#pagination-compact) p,
    [data-testid="column"]:has(#pagination-compact) ~ [data-testid="column"] p {
        font-size: 0.7rem !important; margin: 0 !important;
    }
</style>
"""


def init_session_state():
    if "test_cases" not in st.session_state:
        st.session_state.test_cases = []
    if "test_points" not in st.session_state:
        st.session_state.test_points = []
    if "current_page" not in st.session_state:
        st.session_state.current_page = 1
    if "per_page" not in st.session_state:
        st.session_state.per_page = PER_PAGE_DEFAULT
    if "test_cases_backup" not in st.session_state:
        st.session_state.test_cases_backup = None
    if "editing_case_id" not in st.session_state:
        st.session_state.editing_case_id = None
    if "saved_message" not in st.session_state:
        st.session_state.saved_message = None


def render_pagination(total: int, current_page: int, per_page: int, max_page: int) -> None:
    """
    分页组件：共 X 条 → 上一页 / 页码 / 下一页 → 每页条数 → 跳至 X 页。
    页码为 1-based，布局参考 [2, 4, 1, 2, 2]。
    """
    st.markdown("---")
    col1, col2, col3, col4, col5 = st.columns([2, 4, 1, 2, 2])

    with col1:
        st.markdown(f"共 **{total}** 条")

    with col2:
        col_prev, col_pages, col_next = st.columns([0.5, 9, 0.5])
        with col_prev:
            if st.button("<", key="p_prev", disabled=(current_page == 1), help="上一页", use_container_width=True):
                st.session_state.current_page = current_page - 1
                st.session_state.editing_case_id = None
                st.rerun()
        with col_pages:
            if max_page <= 5:
                page_numbers = list(range(1, max_page + 1))
            else:
                if current_page <= 3:
                    page_numbers = [1, 2, 3, 4, 5, "···", max_page]
                elif current_page >= max_page - 2:
                    page_numbers = [1, "···", max_page - 4, max_page - 3, max_page - 2, max_page - 1, max_page]
                else:
                    page_numbers = [1, "···", current_page - 1, current_page, current_page + 1, "···", max_page]
            page_cols = st.columns(len(page_numbers))
            for idx, p in enumerate(page_numbers):
                with page_cols[idx]:
                    if p == "···":
                        st.markdown("···")
                    else:
                        if st.button(
                            str(p),
                            key=f"p_{p}",
                            type="primary" if p == current_page else "secondary",
                            use_container_width=True,
                        ):
                            if p != current_page:
                                st.session_state.current_page = p
                                st.session_state.editing_case_id = None
                                st.rerun()
        with col_next:
            if st.button(">", key="p_next", disabled=(current_page >= max_page), help="下一页", use_container_width=True):
                st.session_state.current_page = current_page + 1
                st.session_state.editing_case_id = None
                st.rerun()

    with col3:
        cur_idx = PER_PAGE_OPTIONS.index(per_page) if per_page in PER_PAGE_OPTIONS else 0
        new_per = st.selectbox(
            "每页条数",
            options=PER_PAGE_OPTIONS,
            index=cur_idx,
            format_func=lambda x: f"{x}条/页",
            label_visibility="collapsed",
            key="per_page_select",
        )
        if new_per != per_page:
            st.session_state.per_page = new_per
            st.session_state.current_page = 1
            st.session_state.editing_case_id = None
            st.rerun()

    with col4:
        jump_page = st.number_input(
            "",
            min_value=1,
            max_value=max(1, max_page),
            value=current_page,
            step=1,
            label_visibility="collapsed",
            key="jump_num",
        )

    with col5:
        if st.button("跳至", key="jump_btn", use_container_width=True):
            to_page = max(1, min(int(jump_page), max_page))
            st.session_state.current_page = to_page
            st.session_state.editing_case_id = None
            st.rerun()


def handle_pagination(total: int) -> None:
    """统一分页逻辑：current_page 为 1-based，渲染分页组件。"""
    per_page = st.session_state.get("per_page", PER_PAGE_DEFAULT)
    if per_page not in PER_PAGE_OPTIONS:
        per_page = PER_PAGE_DEFAULT
        st.session_state.per_page = per_page
    max_page = max(1, (total + per_page - 1) // per_page)
    current_page = st.session_state.get("current_page", 1)
    current_page = max(1, min(current_page, max_page))
    st.session_state.current_page = current_page
    render_pagination(total, current_page, per_page, max_page)


def run():
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
    st.markdown("""
        <div class="gradient-header">
            <h1>🧪 TestGen AI Testcases </h1>
            <p>基于 AI 自动化生成测试用例。支持文本输入、文档上传、网页链接、图片识别等方式，先解析测试点、按测试点生成用例 · 覆盖更全、列表内直接编辑</p>
        </div>
    """, unsafe_allow_html=True)
    init_session_state()

    client = get_qwen_client()
    if not client:
        st.error("未配置 QWEN_API_KEY，请在 .env 中设置。")
        st.info("示例：QWEN_API_KEY=sk-xxx（阿里云百炼/灵积控制台）")
        return

    # 四种输入方式（含图片识别）
    tab1, tab2, tab3, tab4 = st.tabs(["📝 文本输入", "📎 文档上传", "🔗 网页链接", "🖼️ 图片识别"])
    requirement = ""

    with tab1:
        requirement = st.text_area(
            "需求描述",
            value=st.session_state.get("requirement_text", ""),
            height=200,
            placeholder="请输入需求描述…",
            key="req_text",
        )
        if requirement:
            st.session_state.requirement_text = requirement

    with tab2:
        uploaded = st.file_uploader("支持 PDF、Word(.docx)、TXT、Markdown、XMind(.xmind)、ZIP 压缩包", type=["pdf", "docx", "txt", "md", "xmind", "zip"], key="file_upload")
        if uploaded:
            try:
                requirement = extract_text_from_upload(uploaded)
                st.session_state.requirement_text = requirement
                st.text_area("解析后的文本（可编辑）", value=requirement, height=200, key="req_upload")
            except Exception as e:
                st.error(str(e))

    with tab3:
        url = st.text_input("输入网页 URL", placeholder="https://...", key="url_input")
        if st.button("抓取", key="fetch_btn"):
            if not url or not url.strip():
                st.warning("请输入有效 URL")
            else:
                with st.spinner("正在抓取…"):
                    try:
                        requirement = fetch_url_text(url)
                        st.session_state.requirement_text = requirement
                        st.session_state.req_url = requirement
                        st.text_area("抓取到的正文（可编辑）", value=requirement, height=200, key="req_url")
                    except Exception as e:
                        st.error(str(e))

    with tab4:
        uploaded_image = st.file_uploader("选择图片", type=["png", "jpg", "jpeg", "webp"], key="image_upload")
        # 上传成功后自动解析，不展示图片、不要求点击“识别”按钮
        if uploaded_image:
            last_name = st.session_state.get("last_image_name", "")
            if last_name != uploaded_image.name:
                with st.spinner("🤖 正在识别图片中的文字和需求…"):
                    try:
                        requirement = extract_text_from_image(uploaded_image)
                        st.session_state.requirement_text = requirement
                        st.session_state.req_image = requirement
                        st.session_state.last_image_name = uploaded_image.name
                        st.rerun()
                    except Exception as e:
                        st.error(str(e))
                        st.session_state.last_image_name = None
        if st.session_state.get("req_image"):
            st.text_area("识别的文本（可编辑）", value=st.session_state.get("req_image", ""), height=200, key="req_image")

    requirement = (
        (st.session_state.get("req_text") or "").strip()
        or (st.session_state.get("req_upload") or "").strip()
        or (st.session_state.get("req_url") or "").strip()
        or (st.session_state.get("req_image") or "").strip()
        or (st.session_state.get("requirement_text") or "").strip()
    )
    # 用于智能路由展示：图片 > URL > 文档 > 文本
    if (st.session_state.get("req_image") or "").strip():
        input_source = "image"
    elif (st.session_state.get("req_url") or "").strip():
        input_source = "url"
    elif (st.session_state.get("req_upload") or "").strip():
        input_source = "document"
    else:
        input_source = "text"

    col_gen, col_regen = st.columns(2)
    with col_gen:
        generate_clicked = st.button("生成测试用例", type="primary", use_container_width=True)
    with col_regen:
        regen_clicked = st.button("重新生成", use_container_width=True, help="保留当前需求，重新执行双阶段生成")

    if generate_clicked or regen_clicked:
        if not requirement or not requirement.strip():
            st.warning("请先通过文本/上传/链接/图片方式提供需求内容。")
            st.stop()
        st.session_state.review_result = None
        content = requirement.strip()

        _, _, display_name, desc = get_models_for_input(input_source)
        st.info(f"**{display_name}** | {desc}")

        timer_placeholder = st.empty()
        start_time = time.time()

        try:
            with st.spinner("第1阶段：分析需求，识别测试点…"):
                try:
                    test_points = run_stage1_test_points(client, content)
                except Exception as e:
                    if _is_api_key_error(e):
                        st.error("**API 密钥无效**（401）。请检查项目根目录 `.env` 中的 `QWEN_API_KEY` 是否正确、是否有多余空格，或是否已在阿里云控制台重新生成。详见：[API Key 错误说明](https://help.aliyun.com/zh/model-studio/error-code#apikey-error)")
                    else:
                        st.error(f"测试点分析失败：{str(e)}")
                    st.stop()
            if not test_points:
                st.warning("未识别到测试点，请检查需求描述或重试。")
                st.stop()
            st.session_state.test_points = test_points

            # 主线程内更新耗时（Streamlit 禁止在子线程中调用 UI）
            elapsed1 = time.time() - start_time
            h1, m1 = int(elapsed1 // 3600), int((elapsed1 % 3600) // 60)
            s1 = int(elapsed1 % 60)
            timer_placeholder.caption(f"⏱️ 第1阶段完成，已耗时：{h1:02d}:{m1:02d}:{s1:02d}")

            with st.spinner(f"第2阶段：基于 {len(test_points)} 个测试点生成用例…"):
                try:
                    cases, coverage_map = run_stage2_cases(client, test_points, content)
                except Exception as e:
                    if _is_api_key_error(e):
                        st.error("**API 密钥无效**（401）。请检查 `.env` 中的 `QWEN_API_KEY` 是否正确。详见：[API Key 错误说明](https://help.aliyun.com/zh/model-studio/error-code#apikey-error)")
                    else:
                        st.error(f"用例生成失败：{str(e)}")
                    st.stop()
            if not cases:
                st.warning(
                    "未能解析出有效测试用例，请重试。若测试点较多已自动分批生成仍失败，可尝试：精简上方「解析后的文本」仅保留核心需求后再生成，或换用「文本输入」粘贴部分内容重试。"
                )
                st.stop()

            missing_count = sum(1 for i in range(len(test_points)) if not coverage_map.get(i, False))
            if missing_count > 0:
                cases = fill_missing_cases(client, test_points, content, cases, coverage_map)
                cases = post_process(cases, test_points)

            cases, is_valid = strict_validate_cases(cases, test_points)
            if not is_valid:
                st.error("补全后仍不足，请简化需求或重试。")
                st.stop()

            total_elapsed = time.time() - start_time
            th = int(total_elapsed // 3600)
            tm = int((total_elapsed % 3600) // 60)
            ts = int(total_elapsed % 60)
            st.session_state.test_cases = cases
            st.session_state.current_page = 1
            st.success(f"✅ 共识别 {len(test_points)} 个测试点，生成 {len(cases)} 条用例。总耗时：{th:02d}:{tm:02d}:{ts:02d}")
            st.rerun()
        except Exception as e:
            st.error(f"生成过程出错：{str(e)}")
            st.stop()

    # 结果区
    cases = st.session_state.get("test_cases", [])
    test_points = st.session_state.get("test_points", [])
    if not cases:
        st.info("在上方输入需求并点击「生成测试用例」后，将先分析测试点再生成用例。")
        st.stop()

    # 已保存提示（显示一次后清除）
    if st.session_state.get("saved_message"):
        st.success("✅ 已保存")
        st.session_state.saved_message = None

    # 测试点清单展示（顶部）
    if test_points:
        with st.expander(f"📋 已识别测试点清单（共 {len(test_points)} 个）", expanded=False):
            for i, p in enumerate(test_points[:80], 1):
                st.markdown(f"{i}. {p}")
            if len(test_points) > 80:
                st.caption(f"… 共 {len(test_points)} 个")

    # 分页与表格（无筛选，直接用全部用例）；current_page 为 1-based
    total = len(cases)
    per_page = st.session_state.get("per_page", PER_PAGE_DEFAULT)
    per_page = per_page if per_page in PER_PAGE_OPTIONS else PER_PAGE_DEFAULT
    st.session_state.per_page = per_page
    current_page = st.session_state.get("current_page", 1)
    max_page = max(1, (total + per_page - 1) // per_page)
    current_page = max(1, min(current_page, max_page))
    st.session_state.current_page = current_page
    start = (current_page - 1) * per_page
    end = min(start + per_page, total)
    page_cases = cases[start:end]
    editing_id = st.session_state.get("editing_case_id")

    # 用例列表表格：表头 + 行内编辑（操作列 编辑 / 保存✓ 取消✗）
    TABLE_COL_W = [0.35, 0.4, 1.2, 0.5, 0.45, 1.0, 0.9, 0.9, 1.0, 0.4, 0.35]
    headers = ["序号", "编号", "用例名称", "模块", "类型", "前置条件", "步骤", "测试数据", "预期结果", "优先级", "操作"]
    type_opts = TYPE_OPTS
    pri_opts = ["高", "中", "低"]

    header_cols = st.columns(TABLE_COL_W)
    for i, h in enumerate(headers):
        header_cols[i].markdown(f"**{h}**")

    for row_idx, c in enumerate(page_cases):
        global_idx = start + row_idx
        cid = str(c.get("case_id") or "-")
        is_editing = editing_id == cid
        cols = st.columns(TABLE_COL_W)
        cols[0].write(global_idx + 1)
        cols[1].write(cid)
        if is_editing:
            cols[2].text_input("名称", value=_cell_text(c.get("case_name")), key=f"e_name_{cid}", label_visibility="collapsed")
            cols[3].text_input("模块", value=_get_module(c), key=f"e_mod_{cid}", label_visibility="collapsed")
            tval = _cell_text(c.get("test_type"))
            if tval not in type_opts:
                mapped, _, _ = get_type_tag_style(tval)
                tval = mapped if mapped in type_opts else "功能"
            type_idx = type_opts.index(tval) if tval in type_opts else 0
            cols[4].selectbox("类型", type_opts, index=type_idx, key=f"e_type_{cid}", label_visibility="collapsed")
            cols[5].text_input("前置", value=_cell_text(c.get("precondition")), key=f"e_pre_{cid}", label_visibility="collapsed")
            cols[6].text_area("步骤", value=_cell_text(c.get("steps")), key=f"e_step_{cid}", height=50, label_visibility="collapsed")
            cols[7].text_input("数据", value=_cell_text(c.get("test_data")), key=f"e_data_{cid}", label_visibility="collapsed")
            cols[8].text_area("预期", value=_cell_text(c.get("expected")), key=f"e_exp_{cid}", height=50, label_visibility="collapsed")
            pval = _cell_text(c.get("priority"))
            pri_idx = pri_opts.index(pval) if pval in pri_opts else 1
            cols[9].selectbox("优先级", pri_opts, index=pri_idx, key=f"e_pri_{cid}", label_visibility="collapsed")
            with cols[10]:
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("✓", key=f"save_{cid}", help="保存"):
                        st.session_state.test_cases_backup = [dict(x) for x in st.session_state.test_cases]
                        cases[global_idx].update({
                            "case_name": st.session_state.get(f"e_name_{cid}", c.get("case_name") or ""),
                            "module": st.session_state.get(f"e_mod_{cid}", _get_module(c)),
                            "test_type": st.session_state.get(f"e_type_{cid}", tval),
                            "precondition": st.session_state.get(f"e_pre_{cid}", c.get("precondition") or ""),
                            "steps": st.session_state.get(f"e_step_{cid}", c.get("steps") or ""),
                            "test_data": st.session_state.get(f"e_data_{cid}", c.get("test_data") or ""),
                            "expected": st.session_state.get(f"e_exp_{cid}", c.get("expected") or ""),
                            "priority": st.session_state.get(f"e_pri_{cid}", pval),
                        })
                        st.session_state.editing_case_id = None
                        st.session_state.saved_message = True
                        st.rerun()
                with c2:
                    if st.button("✗", key=f"cancel_{cid}", help="取消"):
                        st.session_state.editing_case_id = None
                        st.rerun()
        else:
            name_display = _cell_text(c.get("case_name"))
            if c.get("_is_placeholder"):
                name_display = "⚠️ " + name_display
            cols[2].write(name_display)
            cols[3].write(_get_module(c))
            t = _cell_text(c.get("test_type"))
            display_text, tag_cls, _ = get_type_tag_style(t)
            cols[4].markdown(f'<span class="type-tag {tag_cls}">{_html_esc(display_text)}</span>', unsafe_allow_html=True)
            cols[5].write(_cell_text(c.get("precondition")))
            cols[6].write(_cell_text(c.get("steps")))
            cols[7].write(_cell_text(c.get("test_data")))
            cols[8].write(_cell_text(c.get("expected")))
            p = _cell_text(c.get("priority"))
            pri_cls = "tag-pri-high" if p == "高" else ("tag-pri-mid" if p == "中" else "tag-pri-low")
            cols[9].markdown(f'<span class="priority-tag {pri_cls}">{_html_esc(p)}</span>', unsafe_allow_html=True)
            with cols[10]:
                if st.button("编辑", key=f"edit_{cid}"):
                    st.session_state.editing_case_id = cid
                    st.rerun()

    # 底部分页：紧凑布局 [共X条] [‹ 页码 ›] [每页条数] [跳至 X 页 跳转]
    handle_pagination(total)

    # 导出（始终为全部用例）
    st.markdown("---")
    st.markdown("### 导出（全部用例）")
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    full_md = f"# 测试用例\n\n共 {len(cases)} 条。\n\n{to_markdown(cases)}"
    with c1:
        st.download_button("📄 Markdown", data=full_md, file_name="test_cases.md", mime="text/markdown", key="dl_md")
    with c2:
        try:
            st.download_button("📊 Excel", data=to_excel_bytes(cases), file_name="test_cases.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key="dl_xlsx")
        except Exception as e:
            st.caption(f"Excel 导出异常：{e}")
    with c3:
        try:
            st.download_button("📝 Word", data=to_word_bytes(cases), file_name="test_cases.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx")
        except Exception as e:
            st.caption(f"Word 导出异常：{e}")
    with c4:
        try:
            st.download_button("📕 PDF", data=to_pdf_bytes(cases), file_name="test_cases.pdf", mime="application/pdf", key="dl_pdf")
        except Exception as e:
            st.caption(f"PDF 导出异常：{e}")
    with c5:
        try:
            st.download_button("🧠 XMind", data=to_xmind_bytes(cases), file_name="test_cases.xmind", mime="application/octet-stream", key="dl_xmind")
            st.caption("若打开时出现修复提示，点「修复并打开」或「关闭」即可正常显示。")
        except Exception as e:
            st.caption(f"XMind 导出异常：{e}")
    with c6:
        try:
            st.download_button("📋 OPML", data=to_opml_bytes(cases), file_name="test_cases.opml", mime="text/x-opml+xml", key="dl_opml")
        except Exception as e:
            st.caption(f"OPML 导出异常：{e}")


if __name__ == "__main__":
    run()
